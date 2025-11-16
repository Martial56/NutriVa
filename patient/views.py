from django.shortcuts import render, redirect, get_object_or_404
#from .forms import PatientForm
from django.db.models import Q
from django.http import HttpResponse, HttpResponseRedirect, HttpResponseBadRequest
from .models import Patient, Constante, Vaccination, Rdv, Nutrition
from datetime import datetime, date
import calendar
import io
from django.template.loader import render_to_string
from django.utils import timezone


try:
    from docx import Document
except:
    Document = None

try:
    import weasyprint
except:
    weasyprint = None


#la vue pour la page d'accueil
def index(request):
    return render(request, 'patient/index.html')

#la vue pour la page de connexion
def login(request):
    return render(request, 'patient/login.html')

#la vue pour la page de creation de patient
def creer_patient(request):
    return render(request, 'patient/creer_patient.html')

#la vue pour la page de saisie des constantes
def constante(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    return render(request, "patient/constante.html", {"patient": patient})

#la vue pour la page rendez-vous
def rdv(request):
    return render(request, 'patient/rdv.html')

#la vue pour la page de saisie des vaccinations
def vaccination(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    return render(request, 'patient/vaccination.html', {"patient": patient})

# la vue pour la page nutrition
def nutrition(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    nutrition = Nutrition.objects.filter(patient=patient).first()
    
    context = {
        "patient": patient,
        "nutrition": nutrition,
    }
    return render(request, 'patient/nutrition.html', context)

#la vue pour la page de liste des patients
def liste_patients(request):
    patients = Patient.objects.all()
    return render(request, "patient/liste_patients.html", {"patients": patients})

# bouton rechercher dans la liste des patients
def rechercher_patients(request):
    query = request.GET.get('search')
    patients = Patient.objects.all()
    if query:
        
        patients = Patient.objects.filter(
            Q(nom__icontains=query) |
            Q(prenom__icontains=query)|
            Q(telephone__icontains=query)
        ).distinct()
    context = {
        "patients": patients,
        "query": query,}
    return render(request, "patient/liste_patients.html", context)

#Bouton enregistrer de la page création de patient
def enregistrement_patient(request):
    if request.method == 'POST':
        nom = request.POST.get('nom')
        prenom = request.POST.get('prenom')
        date_naissance = request.POST.get('date_naissance')
        sexe= request.POST.get('sexe')
        nom_parent = request.POST.get('nom_parent')
        quartier = request.POST.get('quartier')
        telephone = request.POST.get('phone')
        
        date_creation = date.today()
        patient = Patient(
           # code="PT" + str(Patient.objects.count() + 1).zfill(4),
            nom=nom,
            prenom=prenom,
            date_naissance=date_naissance,
            date_creation=date_creation,
            sexe=sexe,
            nom_parent=nom_parent,
            quartier=quartier,
            telephone=telephone
        )
        patient.save()
        return redirect('liste_patients')  # Rediriger vers la liste des patients après l'enregistrement
    return render(request, 'patient/creer_patient.html')

#Générer le code unique pour le patient dans le formulaire de constante
def code_unique(request):
    # Récupère le nombre actuel de patients
    count = Patient.objects.count() + 1

    # Génère le code unique du type NUT0001
    code = "NUT" + str(count).zfill(4)

    # Envoie ce code vers le template
    return render(request, 'patient/nutrition.html', {'code': code})

#Bouton enregistrer de la page saisie des constantes
def enregistrement_constante(request, patient_id):
    if request.method == 'POST':
        date = request.POST.get('date')
        poids = request.POST.get('poids')
        taille = request.POST.get('taille')
        pb = request.POST.get('pb')
        zscore = request.POST.get('zscore')
        #tension = request.POST.get('tension')
        imc = request.POST.get('imc')
        indicec = request.POST.get('indicec')

        patient = get_object_or_404(Patient, id=patient_id)

        Constante.objects.create(
            patient=patient,
            date=date,
            poids=poids,
            taille=taille,
            perimetre_brachial=pb,
            zscore=zscore,
            #tension=tension,
            imc=imc,
            indicecorporel=indicec
        )
        return redirect('rdv')  # Rediriger vers la liste des rendez-vous patients
    return render(request, 'patient/constante.html')

# Fontion pour afficher les rendez-vous des patients
def liste_rdv(request):
    # Récupération des dates depuis le formulaire (méthode GET)
    date_debut_str = request.GET.get('datedebut')
    date_fin_str = request.GET.get('datefin')

    # Valeurs par défaut : aujourd’hui
    today = date.today()

    try:
        date_debut = datetime.strptime(date_debut_str, '%Y-%m-%d').date() if date_debut_str else today
        date_fin = datetime.strptime(date_fin_str, '%Y-%m-%d').date() if date_fin_str else today
    except ValueError:
        date_debut, date_fin = today, today

    # 🔹 Filtrer uniquement les constantes entre les deux dates
    constantes = Constante.objects.filter(date__range=[date_debut, date_fin])

    # 🔹 Extraire uniquement les patients liés à ces constantes
    patients_ids = constantes.values_list('patient_id', flat=True).distinct()
    patients = Patient.objects.filter(id__in=patients_ids)

    # 🔹 Récupérer les nutritions et vaccinations liées à ces patients
    nutritions = Nutrition.objects.filter(patient_id__in=patients_ids, date_visite__range=[date_debut, date_fin])
    vaccinations = Vaccination.objects.filter(patient_id__in=patients_ids, date__range=[date_debut, date_fin])
    rdvs = Rdv.objects.filter(patient_id__in=patients_ids, date_enregistrement__range=[date_debut, date_fin])

    # 🔹 Construire la structure de données à afficher
    data = []
    for patient in patients:
        constante = constantes.filter(patient=patient).first()
        nutrition = nutritions.filter(patient=patient).first()
        vaccination = vaccinations.filter(patient=patient).first()
        rdv = rdvs.filter(patient=patient).first()

        data.append({
            'patient': patient,
            'constante': constante,
            'nutrition': nutrition,
            'vaccination': vaccination,
            'rdv': rdv,
        })

    # 🔹 Envoyer les données au template
    context = {
        'data': data,
        'today': today,
        'date_debut': date_debut,
        'date_fin': date_fin,
    }

    return render(request, 'patient/rdv.html', context)

#enregistrement des vaccinations
def enregistrement_vaccin(request, patient_id):
    if request.method == "POST":
        vaccins = request.POST.getlist('vaccins')  # récupère tous les vaccins cochés
        patient = get_object_or_404(Patient, id=patient_id)
        date_vaccin = date.today()

        #for v in vaccins:
        Vaccination.objects.create(
            patient=patient,
            date=date_vaccin,
            vaccin=vaccins,
        )
        return redirect('rdv')  # ou 'liste_patients', selon ta page de retour

    return render(request, 'patient/vaccination.html')

# enregistrement des infos d'admission en nutrition
def enregistrement_nutrition(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    
    # On vérifie s’il existe déjà une donnée Nutrition pour ce patient
    nutrition = Nutrition.objects.filter(patient=patient).first()

    if request.method == "POST":
        date_admission = request.POST.get("date_admission")
        etat_nutri = request.POST.get("etat_nutri")
        date_sortie = request.POST.get("date_sortie") or None
        motif_sortie = request.POST.get("motif_sortie") or None

        code_unique = f"NUT-{date.today().year}-{str(Nutrition.objects.count() + 1).zfill(4)}",
        date_visite = date.today()
        if nutrition:
            # Mise à jour
            nutrition.code_nutrition = code_unique
            nutrition.date_admission = date_admission
            nutrition.date_visite=date_visite
            nutrition.etat_nutrition = etat_nutri
            nutrition.date_sortie = date_sortie
            nutrition.motif_sortie = motif_sortie
            nutrition.save()
        else:
            # Création
            Nutrition.objects.create(
                patient=patient,
                code_nutrition=code_unique,
                date_admission=date_admission,
                date_visite=date_visite,
                etat_nutrition=etat_nutri,
                date_sortie=date_sortie,
                motif_sortie=motif_sortie,
            )

        return redirect("rdv")  # redirection après enregistrement

    # Préparer les données à envoyer au template
    #context = {
    #    "patient": patient,
     #   "code": nutrition.code_nutrition if nutrition else f"NUT-{patient.id:05d}",
     #   "date_admission": nutrition.date_admission if nutrition else "",
     #   "etat_nutri": nutrition.etat_nutri if nutrition else "",
      #  "date_sortie": nutrition.date_sortie if nutrition else "",
      #  "motif_sortie": nutrition.motif_sortie if nutrition else "",
    #}

    return render(request, "patient/nutrition.html")

def enregistrer_apport_nutrition(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)

    if request.method == "POST":
        depiste = request.POST.get("depiste") or None
        code_depistage = request.POST.get("code_depistage") or None
        resultat = request.POST.get("resultat") or None
        produits = request.POST.getlist("produits") or None  # car c’est une liste de checkboxes

        # Création du rapport
        Rdv.objects.create(
            patient=patient,
            depiste=depiste,
            code_depistage=code_depistage,
            resultat=resultat,
            produits=produits,
        )

        return redirect("rdv")  # après enregistrement

    return render(request, "patient/nutrition.html", {"patient": patient})


# Génération du rapport périodique

def rapports(request):
    # --- Récupération des dates depuis le formulaire ou par défaut sur le mois courant ---
    today = timezone.localdate()
    start_date_str = request.GET.get("date_debut")
    end_date_str = request.GET.get("date_fin")

    if start_date_str and end_date_str:
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        except ValueError:
            return HttpResponseBadRequest("Format de date invalide (utiliser AAAA-MM-JJ)")
    else:
        # par défaut : le mois courant
        start_date = today.replace(day=1)
        last_day = calendar.monthrange(today.year, today.month)[1]
        end_date = today.replace(day=last_day)

    # --- Agrégations ---
    constantes = Constante.objects.filter(date__range=(start_date, end_date)).select_related('patient')
    #patient = Patient.objects.filter(date__range=(start_date, end_date)).select_related('patient')
    vaccinations = Vaccination.objects.filter(date__range=(start_date, end_date)).select_related('patient')
    rdvs = Rdv.objects.filter(date_enregistrement__date__range=(start_date, end_date)).select_related('patient')

    # groupement d’âge
    age_groups = {
        "0-5": (0, 5),
        "6-11": (6, 11),
        "12-23": (12, 23),
        "24-59": (24, 59),
    }
    peses = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}
    total_peses = {"M": 0, "F": 0, "TOTAL": 0}

    for c in constantes:
        age_mois = c.patient.age
        sexe = c.patient.sexe[0].upper() if c.patient.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"
        for g, (low, high) in age_groups.items():
            if low <= age_mois <= high:
                peses[g][sexe] += 1
                peses[g]["TOTAL"] += 1
                total_peses[sexe] += 1
                total_peses["TOTAL"] += 1
                break
    #total_peses = sum(p["TOTAL"] for p in peses.values())
    
     # --- 🧒 Nouveaux enfants (première visite / création du patient) ---
    nouveaux_patients = Patient.objects.filter(date_creation__range=(start_date, end_date))

    premieres_visites = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}
    total_premieres_visites = {"M": 0, "F": 0, "TOTAL": 0}
    for p in nouveaux_patients:
        age_mois = p.age
        sexe = p.sexe[0].upper() if p.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"
        for g, (low, high) in age_groups.items():
            if low <= age_mois <= high:
                premieres_visites[g][sexe] += 1
                premieres_visites[g]["TOTAL"] += 1
                total_premieres_visites[sexe] += 1
                total_premieres_visites["TOTAL"] += 1
                break
    
        # --- Constantes dans la période sélectionnée ---
    constantes = Constante.objects.filter(date__range=(start_date, end_date))


    ### DICTIONNAIRES POUR CHAQUE INDICATEUR ###
    def dict_group():
        return {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}

    obeses = dict_group()
    mam = dict_group()
    mas = dict_group()
    surpoids = dict_group()

    total_obeses = {"M": 0, "F": 0, "TOTAL": 0}
    total_mam = {"M": 0, "F": 0, "TOTAL": 0}
    total_mas = {"M": 0, "F": 0, "TOTAL": 0}
    total_surpoids = {"M": 0, "F": 0, "TOTAL": 0}

    ### 🔍 ANALYSE DES CONSTANTES ###
    for c in constantes:
        patient = c.patient
        age_mois = patient.age

        sexe = patient.sexe[0].upper() if patient.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"

        # Trouver la tranche d'âge
        groupe = None
        for g, (low, high) in age_groups.items():
            if low <= age_mois <= high:
                groupe = g
                break
        if not groupe:
            continue

        categorie = c.indicecorporel.strip().lower()

        # --- Obésité ---
        if categorie == "Obesite" or categorie == "obesite":
            obeses[groupe][sexe] += 1
            obeses[groupe]["TOTAL"] += 1
            total_obeses[sexe] += 1
            total_obeses["TOTAL"] += 1

        # --- MAM ---
        elif categorie == "MAM" or categorie == "mam":
            mam[groupe][sexe] += 1
            mam[groupe]["TOTAL"] += 1
            total_mam[sexe] += 1
            total_mam["TOTAL"] += 1

        # --- MAS ---
        elif categorie == "MAS" or categorie == "mas":
            mas[groupe][sexe] += 1
            mas[groupe]["TOTAL"] += 1
            total_mas[sexe] += 1
            total_mas["TOTAL"] += 1

        # --- Surpoids ---
        elif categorie == "Surpoids" or categorie == "surpoids":
            surpoids[groupe][sexe] += 1
            surpoids[groupe]["TOTAL"] += 1
            total_surpoids[sexe] += 1
            total_surpoids["TOTAL"] += 1

    # --- 🟩 Enfants vus en séance de vaccination PEV (routine) ---
        patient_ids = (
        Vaccination.objects
        .filter(date__range=(start_date, end_date))
        .values_list("patient_id", flat=True)
        .distinct()
    )

    # 2. Charger 1 vaccination par patient
    vaccinations = (
        Vaccination.objects
        .filter(patient_id__in=patient_ids)
        .select_related("patient")
    )

    # Préparation dictionnaires
    pev_routine = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}
    milda = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}
    total_pev_routine = {"M": 0, "F": 0, "TOTAL": 0}
    total_milda = {"M": 0, "F": 0, "TOTAL": 0}

    # Parcours
    for entry in vaccinations:
        patient = entry.patient #patient = Patient.objects.get(id=entry["patient"])
        age_mois = patient.age

        sexe = patient.sexe[0].upper() if patient.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"

        # Trouver la tranche d’âge correspondante
        groupe = None
        for g, (low, high) in age_groups.items():
            if low <= age_mois <= high:
                groupe = g
                break
        if not groupe:
            continue

        # Comptage
        pev_routine[groupe][sexe] += 1
        pev_routine[groupe]["TOTAL"] += 1
        total_pev_routine[sexe] += 1
        total_pev_routine["TOTAL"] += 1
         # ---- 1️⃣ MILDA ----
        if "eabmilda" in entry.vaccin.lower():
            milda[g][sexe] += 1
            milda[g]["TOTAL"] += 1
            total_milda[sexe] += 1
            total_milda["TOTAL"] += 1
        # --- 🟩 Rdv PEV : MILDA, dépistés, dépistés positifs ---

    # 1. Charger les enregistrements RDV dans la période
    rdvs = (
        Rdv.objects
        .filter(date_enregistrement__date__range=(start_date, end_date))
        .select_related("patient")
    )

    # Structures d'âge et sexe
   
    depistes = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}
    positifs = {g: {"M": 0, "F": 0, "TOTAL": 0} for g in age_groups}

    total_depistes = {"M": 0, "F": 0, "TOTAL": 0}
    total_positifs = {"M": 0, "F": 0, "TOTAL": 0}

    for r in rdvs:
        p = r.patient
        age_mois = p.age

        sexe = p.sexe[0].upper() if p.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"

        # Trouver la tranche d'âge
        for g, (low, high) in age_groups.items():
            if low <= age_mois <= high:

                # ---- 2️⃣ Dépistés ----
                if r.depiste == "oui":
                    depistes[g][sexe] += 1
                    depistes[g]["TOTAL"] += 1
                    total_depistes[sexe] += 1
                    total_depistes["TOTAL"] += 1

                # ---- 3️⃣ Dépistés positifs ----
                if r.depiste == "oui" and r.resultat == "positif":
                    positifs[g][sexe] += 1
                    positifs[g]["TOTAL"] += 1
                    total_positifs[sexe] += 1
                    total_positifs["TOTAL"] += 1

                break
  
    # Vaccination
    ANTIGEN_MAP = {
        "BCG": ["bcg"],
       # "Polio 0": ["polio0", "polio 0"],
        "VPO 0": ["VPO0"],
        "VPO 1": ["VPO1"],
        "VPO 2": ["VPO2"],
        "VPO 3": ["VPO3"],
        "VPI 1": ["VPI1"],
        "VPI 2": ["VPI2"],
        "DTC 0": ["Dtc_HepB_Hib_0"],
        "DTC 1": ["Dtc_HepB_Hib_1"],
        "DTC 2": ["Dtc_HepB_Hib_2"],
        "DTC 3": ["Dtc_HepB_Hib_3"],
        "PCV 1": ["PCV1"],
        "PCV 2": ["PCV2"],
        "PCV 3": ["PCV3"],
        "ROTA 1": ["rota1"],
        "ROTA 2": ["rota2"],
        "Rougeole 1": ["RR1"],
        "Rougeole 2": ["RR2"],
        "VAP 1": ["VAP1"],
        "VAP 2": ["VAP2"],
        "VAP 3": ["VAP3"],
        "VAP 4": ["VAP4"],
        "HPV ": ["HPV"],
        "MEN A": ["MenA"],
        "VAA": ["VAA"],
        "Vaccin anti-meningocique": ["vam"],
        "Enfant complètement vacciné": ["vam"],
        "Enfant protégé à la naissance": ["epn"],
        "Enfant ayant bénéficié de MILDA": ["eabmilda"],
        "Vaccin anti-hépatite B": ["ahb"],
        "Vaxigrip": ["Vaxigrip"],
        "VAT 1": ["vat1"],
        "VAT 2": ["vat2"],
        "VAT 1er rappel": ["vatrap1"],
        "VAT 2e rappel": ["vatrap2"],
        "VAT 3e rappel": ["vatrap3"],
    }
    antigen_counts = {k: 0 for k in ANTIGEN_MAP.keys()}
    for v in vaccinations:
        val = v.vaccin.lower()
        for key, patterns in ANTIGEN_MAP.items():
            if any(p in val for p in patterns):
                antigen_counts[key] += 1

    # Produits
    PRODUIT_KEYS = {
        "lait": "Lait",
        "plumpy": "Plumpy Nut",
        "deparasitant": "Déparasitant",
        "vitA100": "Vitamine A100",
        "vitA200": "Vitamine A200",
    }
    produit_counts = {v: 0 for v in PRODUIT_KEYS.values()}
    for r in rdvs:
        for p in (r.produits or []):
            label = PRODUIT_KEYS.get(p)
            if label:
                produit_counts[label] += 1
                    
                    
        # --- RÉCUPÉRATION DES DONNÉES nutrition---
    # Nutrition filtrée sur la date d’admission
    nutritions = (
        Nutrition.objects
        .filter(date_admission__range=(start_date, end_date))
        .select_related("patient")
    )

    # Tranches d'âge (en mois)
    age_groups = {
        "0-5": (0, 5),
        "6-11": (6, 11),
        "12-23": (12, 23),
        "24-59": (24, 59),
        "5-9": (60, 119),
        "10-14": (120, 179),
        "15-19": (180, 239),
        "20-24": (240, 299),
        "25+": (300, 2000),
    }

    # Fonction utilitaire
    def empty():
        return {"M": 0, "F": 0, "TOTAL": 0}

    # Dictionnaires du rapport
    mas_pris = {g: empty() for g in age_groups}
    mam_pris = {g: empty() for g in age_groups}

    mas_gueris = {g: empty() for g in age_groups}
    mam_gueris = {g: empty() for g in age_groups}

    abandon = {g: empty() for g in age_groups}
    deces = {g: empty() for g in age_groups}

    # Totaux globaux
    tot_mas_pris = empty()
    tot_mam_pris = empty()
    tot_mas_gueris = empty()
    tot_mam_gueris = empty()
    tot_abandon = empty()
    tot_deces = empty()


    # Parcours des données
    for n in nutritions:
        p = n.patient
        age_m = p.age

        sexe = p.sexe[0].upper() if p.sexe else "M"
        if sexe not in ["M", "F"]:
            sexe = "M"

        # Groupe d'âge
        groupe = None
        for g, (low, high) in age_groups.items():
            if low <= age_m <= high:
                groupe = g
                break
        if not groupe:
            continue

        # Normalisation
        etat = (n.etat_nutrition or "").lower().strip()
        motif = (n.motif_sortie or "").lower().strip()
        date_sortie = n.date_sortie

        # ==========================
        #       PRISE EN CHARGE
        # ==========================

        # MAS
        if etat == "severe_sans_comp":
            mas_pris[groupe][sexe] += 1
            mas_pris[groupe]["TOTAL"] += 1
            tot_mas_pris[sexe] += 1
            tot_mas_pris["TOTAL"] += 1

        # MAM
        if etat == "modere":
            mam_pris[groupe][sexe] += 1
            mam_pris[groupe]["TOTAL"] += 1
            tot_mam_pris[sexe] += 1
            tot_mam_pris["TOTAL"] += 1


        # ==========================
        #        SORTIES
        # ==========================

        # On doit vérifier la date_sortie !
        if not date_sortie:
            continue

        if not (start_date <= date_sortie <= end_date):
            continue

        # --- Guéris MAS ---
        if etat == "severe_sans_comp" and motif == "gueris":
            mas_gueris[groupe][sexe] += 1
            mas_gueris[groupe]["TOTAL"] += 1
            tot_mas_gueris[sexe] += 1
            tot_mas_gueris["TOTAL"] += 1

        # --- Guéris MAM ---
        if etat == "modere" and motif == "gueris":
            mam_gueris[groupe][sexe] += 1
            mam_gueris[groupe]["TOTAL"] += 1
            tot_mam_gueris[sexe] += 1
            tot_mam_gueris["TOTAL"] += 1

        # --- Abandon ---
        if motif == "abandon":
            abandon[groupe][sexe] += 1
            abandon[groupe]["TOTAL"] += 1
            tot_abandon[sexe] += 1
            tot_abandon["TOTAL"] += 1

        # --- Décès ---
        if motif == "deces":
            deces[groupe][sexe] += 1
            deces[groupe]["TOTAL"] += 1
            tot_deces[sexe] += 1
            tot_deces["TOTAL"] += 1


    # --- 🟩 Enfants ayant reçu du lait ---
    age_groups = {
        "0-5": (0, 5),
        "6-11": (6, 11),
        "12-23": (12, 23),
        "24-59": (24, 59),
    }

    def empty():
        return {"F": 0, "M": 0, "TOTAL": 0}

    lait = {g: empty() for g in age_groups}
    lait["TOTAL"] = empty()

    lait_nouveaux = {g: empty() for g in age_groups}
    lait_nouveaux["TOTAL"] = empty()

    # Récupération des RDV dans la période
    rdvs = Rdv.objects.filter(
        date_enregistrement__range=(start_date, end_date)
    ).select_related("patient")

    # Filtrer lait (compatible SQLite)
    rdvs_lait = [r for r in rdvs if "lait" in r.produits]

    # ───────────────────────────────────────
    # 1. Enfants ayant reçu du lait
    # ───────────────────────────────────────
    for r in rdvs_lait:
        p = r.patient
        age_m = p.age
        sexe = p.sexe.upper()[0] if p.sexe else "M"

        for grp, (low, high) in age_groups.items():
            if low <= age_m <= high:
                lait[grp][sexe] += 1
                lait[grp]["TOTAL"] += 1

                lait["TOTAL"][sexe] += 1
                lait["TOTAL"]["TOTAL"] += 1
                break

    # ───────────────────────────────────────
    # 2. Nouveaux enfants
    # ───────────────────────────────────────
    nouveaux_ids = list(
        Patient.objects.filter(
            date_creation__range=(start_date, end_date)
        ).values_list("id", flat=True)
    )

    # Filtrer dans rdvs_lait (LISTE ➜ list comprehension obligatoire)
    rdvs_nouveaux_lait = [r for r in rdvs_lait if r.patient_id in nouveaux_ids]

    for r in rdvs_nouveaux_lait:
        p = r.patient
        age_m = p.age
        sexe = p.sexe.upper()[0] if p.sexe else "M"

        for grp, (low, high) in age_groups.items():
            if low <= age_m <= high:
                lait_nouveaux[grp][sexe] += 1
                lait_nouveaux[grp]["TOTAL"] += 1

                lait_nouveaux["TOTAL"][sexe] += 1
                lait_nouveaux["TOTAL"]["TOTAL"] += 1
                break

    context = {
        #============ rapport: tableau : Evaluation nutritionnelle ===============
        "date_debut": start_date,
        "date_fin": end_date,
        "peses": peses,
        "total_peses": total_peses,
        "premieres_visites": premieres_visites,
        "total_premieres_visites": total_premieres_visites,
        "obeses": obeses,
        "total_obeses": total_obeses,
        "mam": mam,
        "total_mam": total_mam,
        "mas": mas,
        "total_mas": total_mas,
        "surpoids": surpoids,
        "total_surpoids": total_surpoids,
            ############ rapport: tableau : Vaccinations ############
        "vaccinations": vaccinations,
        "antigen_counts": antigen_counts,
        "produit_counts": produit_counts,
        "rdv_count": rdvs.count(),
        
         ############ rapport: tableau : PEV de routine ############
        "pev_routine": pev_routine,
        "total_pev_routine": total_pev_routine,
        "milda": milda,
        "total_milda": total_milda,
        "depistes": depistes,
        "total_depistes": total_depistes,
        "positifs": positifs,
        "total_positifs": total_positifs,
        "age_groups": age_groups,
        
        ############ rapport: tableau : Nutrition ############
        "mas_pris": mas_pris,
        "tot_mas_pris": tot_mas_pris,
        "mam_pris": mam_pris,
        "tot_mam_pris": tot_mam_pris,
        "mas_gueris": mas_gueris,
        "tot_mas_gueris": tot_mas_gueris,
        "mam_gueris": mam_gueris,
        "tot_mam_gueris": tot_mam_gueris,
        "abandon": abandon,
        "tot_abandon": tot_abandon,
        "deces": deces,
        "tot_deces": tot_deces,
        ############ rapport: tableau : Enfants ayant reçu du lait ############
        "lait": lait,
        "lait_nouveaux": lait_nouveaux,
        
    }

    fmt = request.GET.get("format")
    if fmt == "docx":
        return _report_to_docx(context)
    elif fmt == "pdf":
        return _report_to_pdf(request, context)
    elif fmt == "html_download":
        html = render_to_string("patient/rapport.html", context)
        response = HttpResponse(html, content_type="text/html")
        response["Content-Disposition"] = f'attachment; filename="rapport_{start_date}_{end_date}.html"'
        return response

    return render(request, "patient/rapport.html", context)


# Génération DOCX
def _report_to_docx(context):
    if Document is None:
        return HttpResponseBadRequest("python-docx n'est pas installé.")
    doc = Document()
    doc.add_heading(f"Rapport du {context['date_debut']} au {context['date_fin']}", level=1)

    doc.add_heading("Séances de pesée", level=2)
    t = doc.add_table(rows=1, cols=4)
    hdr = t.rows[0].cells
    hdr[0].text = "Tranche"
    hdr[1].text = "M"
    hdr[2].text = "F"
    hdr[3].text = "Total"
    for g, c in context["peses"].items():
        row = t.add_row().cells
        row[0].text = g
        row[1].text = str(c["M"])
        row[2].text = str(c["F"])
        row[3].text = str(c["TOTAL"])

    doc.add_heading("Vaccinations", level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "Antigène"
    t2.rows[0].cells[1].text = "Nombre"
    for ag, n in context["antigen_counts"].items():
        r = t2.add_row().cells
        r[0].text = ag
        r[1].text = str(n)

    doc.add_heading("Produits distribués", level=2)
    t3 = doc.add_table(rows=1, cols=2)
    t3.rows[0].cells[0].text = "Produit"
    t3.rows[0].cells[1].text = "Quantité"
    for p, n in context["produit_counts"].items():
        r = t3.add_row().cells
        r[0].text = p
        r[1].text = str(n)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    response = HttpResponse(
        bio.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="rapport_{context["date_debut"]}_{context["date_fin"]}.docx"'
    return response


# Génération PDF
def _report_to_pdf(request, context):
    if weasyprint is None:
        return HttpResponseBadRequest("WeasyPrint non installé.")
    html = render_to_string("patient/rapport.html", context)
    pdf = weasyprint.HTML(string=html).write_pdf()
    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="rapport_{context["date_debut"]}_{context["date_fin"]}.pdf"'
    return response